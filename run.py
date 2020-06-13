from docx import Document
from docx.shared import Inches
import telebot
import urllib
import os
from PIL import Image, ImageDraw
import re

API_TOKEN = "1274266029:AAHhPrpdgzzD-z4sYYoiqXhkswUeyBXe_DE"

bot = telebot.TeleBot(API_TOKEN)

user_dict = {}


class User:
    def __init__(self, name):
        self.name = name
        self.age = None
        self.sex = None


@bot.message_handler(commands=['start'])
def handle_start(message):
    user_markup = telebot.types.ReplyKeyboardMarkup()
    user_markup.row('/Dox')
    user_markup.row('/Exdox')

    bot.send_message(message.from_user.id, "ssss", reply_markup=user_markup)




@bot.message_handler(commands=["Dox"])
def send_welcome(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        global msg
        msg = bot.reply_to(message, 'Спеціальність')

        bot.register_next_step_handler(msg, name_of_specialty)
    except Exception as e:
        bot.reply_to(message, '1')


def name_of_specialty(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'Освітня програма')
        global a1
        a1 = (format(message.text))
        bot.register_next_step_handler(msg, specialization)
    except Exception as e:
        bot.reply_to(message, "5")


def specialization(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'Спеціалізація')
        global a2
        a2 = (format(message.text))
        bot.register_next_step_handler(msg, educational_program)
    except Exception as e:
        bot.reply_to(message, "5")




def educational_program(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'ОКР')
        global a3
        a3 = (format(message.text))
        bot.register_next_step_handler(msg, okp)
    except Exception as e:
        bot.reply_to(message, "5")


def okp(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'назва дисциплінини')
        global a4
        a4 = (format(message.text))
        bot.register_next_step_handler(msg, subjects)
    except Exception as e:
        bot.reply_to(message, "5")


def subjects(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'кількість питань у білеті')
        global a5
        a5 = (format(message.text))
        bot.register_next_step_handler(msg, number_of_questions_in_the_ticket)
    except Exception as e:
        bot.reply_to(message, '6')


def number_of_questions_in_the_ticket(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'номер протоколу')
        global a6
        a6 = (format(message.text))
        bot.register_next_step_handler(msg, meeting_of_the_department)
    except Exception as e:
        bot.reply_to(message, '6')




def meeting_of_the_department(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'засідання кафедри ')
        global a7
        a7 = (format(message.text))
        bot.register_next_step_handler(msg, data)
    except Exception as e:
        bot.reply_to(message, '6')



def data(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        msg = bot.reply_to(message, 'Дата')
        global a8
        a8 = (format(message.text))
        bot.register_next_step_handler(msg, red_dox)
    except Exception as e:
        bot.reply_to(message, '6')












def red_dox(message):
    try:
        global a9
        a9 = (format(message.text))



        chat_id = message.chat.id
        name = message.text

        global xa
        xa = (format(message.text))

        print(msg)

        document = Document()
        p = document.add_heading('КИЇВСЬКИЙ УНІВЕРСИТЕТ імені БОРИСА ГРІНЧЕНКА')
        p = document.add_heading('Факультет інформаційних технологій та управління')

        records = (
            ("Спеціалізація ", a1),
            ("Освітня програма", a2),
            ("Спеціальність ", a3),
            ("ОКР ", a4),
            ("Дисципліна ", a5)
        )

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '1'
        hdr_cells[1].text = '2'

        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id





        document.add_page_break()
        document.save('demo.docx')

    except Exception as e:
        global a11
        a11 = (format(message.text))
        chat_id = message.chat.id
        name = message.text
        global a12
        a12 = (format(message.text))
        print(msg)

        document = Document()
        p = document.add_heading('КИЇВСЬКИЙ УНІВЕРСИТЕТ імені БОРИСА ГРІНЧЕНКА')
        p = document.add_heading('      Факультет інформаційних технологій та управління', 3)




        records = (
            ("Спеціальність ", a1),
            ("Спеціалізація", a2),
            ("Освітня програма  ", a3),
            ("ОКР ", a4 ),
            ("Дисципліна ", a5)
        )

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Категорія'
        hdr_cells[1].text = 'Назва'

        for qty, id in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id



        p = document.add_heading('Комплексний тест з дисципліни (20 балів).', 6)
        p = document.add_heading('Практичне завдання №1 (10 балів).', 6)
        p = document.add_heading('Практичне завдання №2 ( 10 балів).', 6)
        p = document.add_heading('')
        p = document.add_heading('')
        p = document.add_heading('')
        p = document.add_heading('')
        p = document.add_heading('')




        document.add_heading(
            'Розглянуто та затверджено на засіданні кафедри ', level=7
        )
        document.add_heading(
            'комп’ютерних наук та математики ', level=7
        )
        document.add_heading(
            'Протокол №'+a7+'від '+a9, level=7
        )

        document.add_heading(
            'Протокол №'+a7+'від '+a9, level=7
        )

        document.add_heading(
            'Завідувач кафедри Литвин О.С.______Екзаменатор Кучаковська Г.А._______', 9
        )



        document.add_page_break()
        document.save('demo.docx')




@bot.message_handler(commands=["Exdox"])
def send_welcome(message):
    try:
        chat_id = message.chat.id
        name = message.text
        user = User(name)
        user_dict[chat_id] = user
        global msg
        msg = bot.reply_to(message, 'Вкажіть файл')


    except Exception as e:
        print("sdds")








bot.polling()







